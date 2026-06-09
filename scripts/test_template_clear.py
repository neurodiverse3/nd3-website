import docx

def test():
    template_path = r"C:\Users\Ollie\Documents\Projects\ND3 Website\neurodivers3_template_light.docx"
    doc = docx.Document(template_path)
    
    # Remove all body paragraphs
    # Note: iterating directly can skip elements when we remove, so we convert to list
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
        
    # Remove all body tables
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)
        
    # Add a title
    p_title = doc.add_paragraph()
    run = p_title.add_run("Test Title")
    run.bold = True
    run.font.name = "Century Gothic"
    run.font.size = docx.shared.Pt(24)
    
    # Add a paragraph
    p = doc.add_paragraph()
    p.add_run("This is some text using the light template style. The header and footer should be preserved.")
    
    output_path = r"C:\Users\Ollie\Documents\Projects\ND3 Website\scripts\test_output.docx"
    doc.save(output_path)
    print(f"Test saved to: {output_path}")

test()
