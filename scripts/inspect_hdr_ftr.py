import docx

def inspect_header_footer(path):
    doc = docx.Document(path)
    section = doc.sections[0]
    
    print(f"\n--- Header of {path} ---")
    header = section.header
    print(f"Header has {len(header.paragraphs)} paragraphs.")
    for i, p in enumerate(header.paragraphs):
        print(f"p[{i}]: text='{p.text}'")
    print(f"Header has {len(header.tables)} tables.")
    for i, t in enumerate(header.tables):
        print(f"Table {i} rows: {len(t.rows)}, cols: {len(t.columns)}")
        for ri, r in enumerate(t.rows):
            for ci, c in enumerate(r.cells):
                print(f"  Cell [{ri},{ci}]: text='{c.text.strip()}'")

    print(f"\n--- Footer of {path} ---")
    footer = section.footer
    print(f"Footer has {len(footer.paragraphs)} paragraphs.")
    for i, p in enumerate(footer.paragraphs):
        print(f"p[{i}]: text='{p.text}'")
    print(f"Footer has {len(footer.tables)} tables.")
    for i, t in enumerate(footer.tables):
        print(f"Table {i} rows: {len(t.rows)}, cols: {len(t.columns)}")
        for ri, r in enumerate(t.rows):
            for ci, c in enumerate(r.cells):
                print(f"  Cell [{ri},{ci}]: text='{c.text.strip()}'")

inspect_header_footer(r"C:\Users\Ollie\Documents\Projects\ND3 Website\neurodivers3_template_light.docx")
