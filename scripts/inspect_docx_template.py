import docx

def inspect_template(path):
    doc = docx.Document(path)
    print(f"\n--- Inspecting template: {path} ---")
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    for i, p in enumerate(doc.paragraphs[:10]):
        print(f"p[{i}]: text='{p.text}', style='{p.style.name}'")
    print(f"Number of tables: {len(doc.tables)}")
    print(f"Number of sections: {len(doc.sections)}")

inspect_template(r"C:\Users\Ollie\Documents\Projects\ND3 Website\neurodivers3_template_light.docx")
inspect_template(r"C:\Users\Ollie\Documents\Projects\ND3 Website\neurodivers3_template_dark.docx")
